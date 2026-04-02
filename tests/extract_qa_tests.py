#!/usr/bin/env python3
"""Extract all test cases from ECOM360_QA_TEST_PLAN.docx into JSON for automated execution."""
import json
from docx import Document

doc = Document('/Users/surenderaggarwal/Projects/ecom360/docs/ECOM360_QA_TEST_PLAN.docx')

sections = []
current_section = None
current_subsection = None

# Walk through document elements in order
for elem in doc.element.body:
    tag = elem.tag.split('}')[-1]
    if tag == 'p':
        from docx.oxml.ns import qn
        text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
        text = text.strip()
        if not text:
            continue
        # Detect section headers
        if text.startswith('SECTION '):
            current_section = {'title': text, 'subsections': [], 'description': ''}
            sections.append(current_section)
        elif current_section:
            pPr = elem.find(qn('w:pPr'))
            style = ''
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style = pStyle.get(qn('w:val'), '')
            if 'Heading2' in style or 'Heading 2' in style:
                current_subsection = {'title': text, 'tests': []}
                current_section['subsections'].append(current_subsection)
            elif not current_section.get('subsections'):
                current_section['description'] = text

    elif tag == 'tbl' and current_subsection is not None:
        # Find this table in doc.tables
        for t in doc.tables:
            if t._element is elem:
                headers = [c.text.strip().lower() for c in t.rows[0].cells]
                if 'test id' not in headers:
                    break
                for row in t.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    test = {}
                    for i, h in enumerate(headers):
                        if i < len(cells):
                            test[h] = cells[i]
                    if test.get('test id'):
                        current_subsection['tests'].append(test)
                break

# Print summary
total = 0
for s in sections:
    count = sum(len(sub['tests']) for sub in s['subsections'])
    total += count
    print(f"{s['title']}: {count} tests across {len(s['subsections'])} subsections")
    for sub in s['subsections']:
        if sub['tests']:
            print(f"  {sub['title']}: {len(sub['tests'])} tests")

print(f"\nTOTAL: {total} test cases")

# Save to JSON
with open('/Users/surenderaggarwal/Projects/ecom360/tests/qa_test_cases.json', 'w') as f:
    json.dump(sections, f, indent=2)
print("\nSaved to tests/qa_test_cases.json")
